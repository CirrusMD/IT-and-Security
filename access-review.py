import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def main():
    # Initialize a Document
    doc = Document()
    doc.add_heading('Access Review Report', 0)

    # Add a date
    doc.add_paragraph('Date: Some Date Here')

    # Load data
    user_list_df = pd.read_csv('/mnt/data/jcuserlist_20230925.csv')
    user_groups_df = pd.read_csv('/mnt/data/users-to-user-groups-6511da5f8d205ef0d9bcdbe7-20230925Z.csv')
    user_apps_df = pd.read_csv('/mnt/data/users-to-sso-applications-6511daca2624fcca7fea9f4c-20230925Z.csv')

    # Step 1: User Status - Identify inactive or suspicious accounts
    inactive_accounts = user_list_df[user_list_df['state'] == 'INACTIVE']
    suspended_accounts = user_list_df[user_list_df['state'] == 'SUSPENDED']
    doc.add_heading('1. User Status: Inactive or Suspicious Accounts', level=1)
    doc.add_paragraph(f'Number of Inactive Accounts: {len(inactive_accounts)}')
    doc.add_paragraph(f'Number of Suspended Accounts: {len(suspended_accounts)}')

    # Generate a bar plot for User Status
    sns.set(style="whitegrid")
    plt.figure(figsize=(10, 6))
    sns.barplot(x=['Inactive', 'Suspended'], y=[len(inactive_accounts), len(suspended_accounts)])
    plt.title('User Status Summary')
    plt.xlabel('Account Status')
    plt.ylabel('Number of Accounts')
    plt.savefig('user_status_summary.png')
    doc.add_picture('user_status_summary.png', width=Inches(4.0))

    # Step 2: Group Membership
    group_counts = user_groups_df['user_group_name'].value_counts()
    doc.add_heading('2. Group Membership', level=1)
    doc.add_paragraph(f'Top 5 Groups with Most Members:')
    doc.add_paragraph(f'{group_counts.head()}')

    # Step 3: Application Access
    app_access_counts = user_apps_df['application_name'].value_counts()
    doc.add_heading('3. Application Access', level=1)
    doc.add_paragraph(f'Top 5 Most Accessed Applications:')
    doc.add_paragraph(f'{app_access_counts.head()}')

    # Step 4: Data Quality
    missing_values = user_list_df.isnull().sum()
    doc.add_heading('4. Data Quality', level=1)
    doc.add_paragraph(f'Missing Values in User List:')
    doc.add_paragraph(f'{missing_values}')

    # Step 5: High-Risk Combinations
    admin_groups = user_groups_df[user_groups_df['user_group_name'].str.contains('Admin')]
    sensitive_apps = user_apps_df['application_name'].value_counts().index[:5]  # Top 5 as example
    high_risk_combinations = pd.merge(admin_groups, user_apps_df, on=['username', 'email'])
    high_risk_combinations = high_risk_combinations[high_risk_combinations['application_name'].isin(sensitive_apps)]
    doc.add_heading('5. High-Risk Combinations', level=1)
    doc.add_paragraph(f'High-Risk Combinations:')
    doc.add_paragraph(f'{high_risk_combinations.head()}')


export_csv_option = input("Do you want to export additional CSV files? (y/n): ")
if export_csv_option.lower() == 'y':
    print("Choose which additional CSV files you want to export:")
    print("1: List of users and the number of groups assigned, and another column that lists all the groups")
    print("2: List of Groups and the total number of users")
    print("3: Inactive Users With Group Access")
    print("4: Groups With Sensitive App Access")
    print("5: Users Without Two-Factor Authentication")
    print("6: Rarely Used Apps")
    print("7: Users With Admin Privileges")
    print("8: Data Completeness Report")
    print("9: Expired Account Credentials")
    selected_options = input("Enter the numbers of the options you want, separated by commas (e.g., 1,3,5): ").split(",")
    
    if '1' in selected_options:
        # 1) List of users and the number of groups assigned, and another column that lists all the groups
        user_group_count = user_groups_df.groupby('username')['user_group_name'].agg(['count', lambda x: ', '.join(x)]).reset_index()
        user_group_count.columns = ['Username', 'Number_of_Groups', 'List_of_Groups']
        user_group_count.to_csv('User_Group_Info.csv', index=False)
    
    if '2' in selected_options:
        # 2) List of Groups and the total number of users
        group_user_count = user_groups_df.groupby('user_group_name')['username'].count().reset_index()
        group_user_count.columns = ['Group_Name', 'Number_of_Users']
        group_user_count.to_csv('Group_User_Info.csv', index=False)    

    if '3' in selected_options:
        # 3) Inactive Users With Group Access
        inactive_users_with_groups.to_csv('Inactive_Users_With_Group_Access.csv', index=False)
        
    if '4' in selected_options:
        # 4) Groups With Sensitive App Access
        groups_with_sensitive_apps.to_csv('Groups_With_Sensitive_App_Access.csv', index=False)
        
    if '5' in selected_options:
        # 5) Users Without Two-Factor Authentication
        user_list_df.to_csv('Users_Without_Two_Factor_Authentication.csv', index=False)
        
    if '6' in selected_options:
        # 6) Rarely Used Apps
        rarely_used_apps.to_csv('Rarely_Used_Apps.csv', index=False)
        
    if '7' in selected_options:
        # 7) Users With Admin Privileges
        users_with_admin_privileges.to_csv('Users_With_Admin_Privileges.csv', index=False)
        
    if '8' in selected_options:
        # 8) Data Completeness Report
        data_completeness_report.to_csv('Data_Completeness_Report.csv', index=False)
        
    if '9' in selected_options:
        # 9) Expired Account Credentials
        expired_account_credentials.to_csv('Expired_Account_Credentials.csv', index=False)
        
    print("Selected CSV files have been exported.")        
    
    # Save the document
    current_time = datetime.now().strftime('%Y%m%d%H%M%S')
    doc.save(f'Access_Review_Report_{current_time}.docx')

if __name__ == '__main__':
    main()
